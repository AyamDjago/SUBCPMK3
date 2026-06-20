"""
Generate artikel_cbr.docx - Format mendekati JOIV Template
Universitas Muhammadiyah Malang - Penalaran Komputer
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

# ==============================================================
# Konfigurasi Warna & Font (JOIV-style)
# ==============================================================
FONT_BODY = "Times New Roman"
FONT_TITLE = "Times New Roman"

def set_font(run, name=FONT_BODY, size_pt=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=4):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(12)

def add_heading(doc, text, level=1, numbered=True):
    """Add section heading"""
    para = doc.add_paragraph()
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT,
             space_before=6, space_after=4)
    run = para.add_run(text.upper() if level == 1 else text)
    size = 10 if level >= 1 else 14
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    run.bold = True
    return para

def add_body_para(doc, text, italic=False):
    """Add body paragraph"""
    para = doc.add_paragraph()
    para_fmt(para)
    if text:
        run = para.add_run(text)
        set_font(run, italic=italic)
    return para

def add_table(doc, headers, rows, caption, table_num):
    """Add formatted table with caption"""
    caption_para = doc.add_paragraph()
    para_fmt(caption_para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    caption_run = caption_para.add_run(f"Tabel {table_num}. {caption}")
    caption_run.font.name = FONT_BODY
    caption_run.font.size = Pt(10)
    caption_run.bold = True

    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            run.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Gray background for header
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx+1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = str(cell_text)
            for run in row[c_idx].paragraphs[0].runs:
                run.font.name = FONT_BODY
                run.font.size = Pt(10)
            row[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # Space after table
    return table


# ==============================================================
# BUAT DOKUMEN
# ==============================================================
doc = Document()

# --- Halaman & Margin (A4, margin: T/B=3cm, L/R=2cm) ---
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ==============================================================
# JUDUL
# ==============================================================
title_para = doc.add_paragraph()
title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(0)
title_para.paragraph_format.space_after = Pt(6)
title_run = title_para.add_run(
    "Penerapan Case-Based Reasoning Berbasis IndoBERT\n"
    "untuk Analisis Putusan Pidana Pemalsuan"
)
title_run.font.name = FONT_TITLE
title_run.font.size = Pt(18)
title_run.bold = True

# ==============================================================
# PENULIS
# ==============================================================
author_para = doc.add_paragraph()
author_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_para.paragraph_format.space_after = Pt(2)
ar = author_para.add_run("[NAMA LENGKAP ANDA]")
ar.font.name = FONT_BODY
ar.font.size = Pt(11)
ar.bold = True

affil_para = doc.add_paragraph()
affil_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil_para.paragraph_format.space_after = Pt(2)
affr = affil_para.add_run(
    "Teknik Informatika, Fakultas Teknik, Universitas Muhammadiyah Malang\n"
    "Jl. Raya Tlogomas No. 246, Malang, Jawa Timur, Indonesia\n"
    "e-mail: [EMAIL@webmail.umm.ac.id]"
)
affr.font.name = FONT_BODY
affr.font.size = Pt(10)

doc.add_paragraph()

# ==============================================================
# ABSTRACT
# ==============================================================
add_heading(doc, "Abstract")

abstract_text = (
    "Case-Based Reasoning (CBR) is a problem-solving method that utilizes past experiences to solve new problems. "
    "This study implements a CBR system to analyze criminal forgery court decisions from the Supreme Court of the "
    "Republic of Indonesia. The system processes 35 court decision documents through five main stages: case base "
    "construction, case representation, case retrieval using IndoBERT embeddings, solution reuse with weighted "
    "similarity, and evaluation. The IndoBERT language model (indobenchmark/indobert-base-p1) is employed to generate "
    "768-dimensional semantic embeddings from case documents, which are then compared using cosine similarity for case "
    "retrieval. The solution reuse stage applies a weighted similarity voting mechanism to predict verdicts for new cases. "
    "Experimental results using an 80:20 train-test split demonstrate a prediction accuracy of 71.43%, with weighted "
    "precision, recall, and F1-score of 0.51, 0.71, and 0.60, respectively. The system achieves perfect recall for "
    "imprisonment verdicts but fails to identify acquittal cases. This study contributes a domain-specific CBR framework "
    "for Indonesian legal text analysis and identifies key challenges in automated verdict prediction from unstructured "
    "court documents."
)
para_abs = doc.add_paragraph()
para_fmt(para_abs)
r_abs = para_abs.add_run(abstract_text)
set_font(r_abs, italic=True)

kw_para = doc.add_paragraph()
para_fmt(kw_para)
kw_label = kw_para.add_run("Kata Kunci")
set_font(kw_label, bold=True, italic=True)
kw_rest = kw_para.add_run(
    " \u2014 Case-Based Reasoning, IndoBERT, Pemalsuan, Cosine Similarity, Prediksi Putusan."
)
set_font(kw_rest, italic=True)

doc.add_paragraph()

# ==============================================================
# I. PENDAHULUAN
# ==============================================================
add_heading(doc, "I. Pendahuluan")

intro_paragraphs = [
    ("Sistem peradilan di Indonesia menghasilkan ribuan putusan pengadilan setiap tahunnya, khususnya dalam bidang pidana umum. "
     "Putusan-putusan tersebut memuat informasi hukum yang sangat berharga, termasuk fakta persidangan, dakwaan, pertimbangan hukum, "
     "dan amar putusan yang dapat dijadikan referensi untuk kasus-kasus serupa di masa mendatang [1]. Namun, volume dokumen putusan "
     "yang sangat besar dan formatnya yang tidak terstruktur menjadikan pencarian dan analisis putusan secara manual sebagai proses "
     "yang memakan waktu dan rentan terhadap inkonsistensi [2]."),

    ("Tindak pidana pemalsuan yang diatur dalam Pasal 263 hingga 276 Kitab Undang-Undang Hukum Pidana (KUHP) merupakan salah satu "
     "jenis perkara yang cukup sering ditangani di pengadilan Indonesia. Perkara-perkara pemalsuan melibatkan berbagai modus operandi, "
     "mulai dari pemalsuan surat, akta otentik, hingga dokumen identitas, sehingga analisis yurisprudensi dalam domain ini memerlukan "
     "pemahaman kontekstual yang mendalam terhadap setiap kasus [3]."),

    ("Case-Based Reasoning (CBR) merupakan pendekatan penalaran berbasis pengalaman yang menyelesaikan masalah baru dengan merujuk "
     "pada solusi dari kasus-kasus serupa di masa lalu [1]. Pendekatan ini sangat relevan untuk domain hukum karena prinsip yurisprudensi "
     "dalam sistem hukum pada dasarnya juga mengandalkan preseden \u2014 yakni putusan-putusan terdahulu yang dijadikan acuan dalam memutus "
     "perkara serupa. Siklus CBR yang dikemukakan oleh Aamodt dan Plaza [1] terdiri dari empat tahap utama: Retrieve (menemukan kasus "
     "yang mirip), Reuse (mengadaptasi solusi), Revise (mengevaluasi solusi), dan Retain (menyimpan pengalaman baru)."),

    ("Perkembangan teknologi Natural Language Processing (NLP) berbasis transformer telah membuka peluang baru dalam pemrosesan teks "
     "hukum. Model bahasa pra-latih seperti BERT [4] telah terbukti efektif dalam menangkap representasi semantik dari teks. Untuk "
     "bahasa Indonesia, model IndoBERT [5] yang dilatih pada korpus bahasa Indonesia berskala besar memberikan kemampuan pemahaman "
     "kontekstual yang superior dibandingkan metode bag-of-words tradisional. Pemanfaatan IndoBERT sebagai mesin embedding dalam sistem "
     "CBR memungkinkan pengukuran kemiripan antar kasus hukum secara semantik, bukan sekadar kecocokan kata kunci [6]."),

    ("Beberapa penelitian terdahulu telah mengeksplorasi penerapan kecerdasan buatan dalam prediksi putusan hukum. Collenette et al. [8] "
     "mengembangkan sistem AI yang dapat menjelaskan penalaran berbasis kasus pada Pengadilan Hak Asasi Manusia Eropa. Cui et al. [9] "
     "melakukan survei komprehensif tentang metode prediksi putusan pengadilan, termasuk pendekatan berbasis deep learning. Di Indonesia, "
     "penelitian tentang penerapan NLP pada dokumen hukum mulai berkembang, dengan studi seperti yang dilakukan oleh Kemala dan Shiddiqi [11] "
     "yang mengklasifikasikan putusan pengadilan pajak, serta Yulianti et al. [12] yang menerapkan model transformer untuk Named Entity "
     "Recognition pada dokumen hukum Indonesia."),

    ("Penelitian ini bertujuan untuk mengimplementasikan sistem CBR yang memanfaatkan model IndoBERT untuk menganalisis putusan pidana "
     "pemalsuan dari Mahkamah Agung Republik Indonesia. Kontribusi utama dari penelitian ini meliputi: (1) pengembangan pipeline CBR "
     "end-to-end untuk dokumen hukum Indonesia, (2) pemanfaatan IndoBERT sebagai mesin representasi kasus yang menghasilkan embedding "
     "semantik berdimensi 768, dan (3) evaluasi komprehensif terhadap performa sistem menggunakan metrik information retrieval dan "
     "klasifikasi standar."),
]

for text in intro_paragraphs:
    add_body_para(doc, text)

# ==============================================================
# II. METODOLOGI
# ==============================================================
add_heading(doc, "II. Metodologi")

add_body_para(doc,
    "Sistem CBR yang dikembangkan dalam penelitian ini mengimplementasikan siklus CBR melalui lima tahap utama: "
    "(1) pembangunan case base, (2) representasi kasus, (3) case retrieval menggunakan IndoBERT, "
    "(4) solution reuse dengan weighted similarity, dan (5) evaluasi."
)

# A. Pembangunan Case Base
add_heading(doc, "A. Pembangunan Case Base", level=2)
add_body_para(doc,
    "Tahap pertama adalah pengumpulan dan pemrosesan dokumen putusan pengadilan dalam format PDF menjadi teks bersih. "
    "Dataset terdiri dari 35 dokumen putusan pidana pemalsuan yang diperoleh dari Direktori Putusan Mahkamah Agung "
    "Republik Indonesia (https://putusan3.mahkamahagung.go.id/). Seluruh putusan merupakan perkara pidana umum "
    "yang berkaitan dengan tindak pidana pemalsuan sebagaimana diatur dalam Pasal 263 sampai dengan Pasal 276 KUHP."
)
add_body_para(doc,
    "Konversi dokumen PDF ke teks dilakukan menggunakan pustaka pdfminer.six dengan fungsi extract_text(). "
    "Teks hasil konversi kemudian melewati pipeline pembersihan tujuh langkah: (1) penghapusan watermark, header, "
    "dan footer menggunakan pola regex; (2) penghapusan nomor halaman; (3) normalisasi karakter Unicode; "
    "(4) penghapusan karakter non-ASCII berlebih; (5) normalisasi spasi; (6) konversi lowercase; dan "
    "(7) penghapusan spasi di awal dan akhir teks."
)
add_body_para(doc,
    "Validasi dilakukan terhadap setiap dokumen dengan kriteria: jumlah kata minimum 200 kata dan rasio kemunculan "
    "kata kunci domain minimal 80% dari lima kata kunci: terdakwa, dakwaan, tuntutan, putusan, dan majelis hakim. "
    "Seluruh 35 dokumen memenuhi kriteria validasi dengan rata-rata 8.662 kata per dokumen."
)

# B. Representasi Kasus
add_heading(doc, "B. Representasi Kasus", level=2)
add_body_para(doc,
    "Tahap kedua mengekstrak metadata dan konten kunci dari setiap dokumen teks menggunakan pola regular expression "
    "(regex). Setiap kasus direpresentasikan dalam struktur data yang mencakup 18 atribut, dikelompokkan menjadi empat "
    "kategori: (1) identitas kasus: case_id, original_filename; (2) metadata hukum: no_perkara, tanggal, jenis_perkara, "
    "pasal, terdakwa, penuntut, pengadilan; (3) konten kunci: ringkasan_fakta, dakwaan, amar_putusan, vonis; dan "
    "(4) fitur statistik: word_count, sentence_count, unique_words, top_keywords, pemalsuan_keywords."
)

# C. Case Retrieval
add_heading(doc, "C. Case Retrieval Menggunakan IndoBERT", level=2)
add_body_para(doc,
    "Tahap ketiga merupakan inti dari proses Retrieve dalam siklus CBR. Model IndoBERT (indobenchmark/indobert-base-p1) [5] "
    "digunakan untuk menghasilkan representasi vektor (embedding) berdimensi 768 dari setiap kasus. Teks embedding "
    "dibangun dengan menggabungkan tiga komponen: pasal yang didakwakan, ringkasan dakwaan (maks. 300 karakter), dan "
    "ringkasan fakta persidangan (maks. 300 karakter), dalam format: \"Pasal: {pasal} Dakwaan: {dakwaan} Fakta: {fakta}\"."
)
add_body_para(doc,
    "Proses pembuatan embedding terdiri dari: (1) tokenisasi menggunakan tokenizer IndoBERT dengan panjang maksimum 512 "
    "token; (2) forward pass melalui model tanpa komputasi gradien; (3) mean pooling atas token embeddings berdasarkan "
    "attention mask; dan (4) normalisasi L2. Data dibagi dengan stratified random split 80:20 (28 training, 7 testing, "
    "random_state=42). Kemiripan antar kasus dihitung menggunakan cosine similarity."
)

# D. Solution Reuse
add_heading(doc, "D. Solution Reuse dengan Weighted Similarity", level=2)
add_body_para(doc,
    "Setelah Top-K kasus terdekat diperoleh (K=5), sistem memprediksi vonis menggunakan weighted similarity voting. "
    "Label vonis dinormalisasi ke lima kategori: BEBAS, PENJARA_X_TAHUN, PENJARA_X_BULAN, PENJARA_LAIN, dan TIDAK_DIKETAHUI. "
    "Metode prediksi utama menjumlahkan skor cosine similarity untuk setiap kategori vonis dan memilih kategori dengan "
    "bobot kumulatif tertinggi."
)

# E. Evaluasi
add_heading(doc, "E. Evaluasi", level=2)
add_body_para(doc,
    "Evaluasi dilakukan menggunakan dua kelompok metrik. Untuk evaluasi retrieval digunakan Precision@K, Recall@K, F1@K, "
    "dan Mean Reciprocal Rank (MRR) pada K \u2208 {1, 3, 5}. Ground truth ditentukan berdasarkan kecocokan pasal yang "
    "didakwakan. Untuk evaluasi prediksi digunakan metrik klasifikasi standar: Accuracy, Precision (weighted), "
    "Recall (weighted), dan F1-Score (weighted)."
)

# ==============================================================
# III. IMPLEMENTASI
# ==============================================================
add_heading(doc, "III. Implementasi")

add_body_para(doc,
    "Implementasi sistem dilakukan menggunakan Python 3.10.11 dalam Jupyter Notebook. Tabel I menyajikan spesifikasi "
    "pustaka utama yang digunakan."
)

add_table(doc,
    headers=["Pustaka", "Versi", "Fungsi"],
    rows=[
        ["transformers", "4.41.2", "Model IndoBERT dan tokenizer"],
        ["torch (PyTorch)", "2.3.0", "Komputasi tensor dan inferensi"],
        ["scikit-learn", "1.4.2", "Pembagian data & metrik evaluasi"],
        ["pdfminer.six", "20260107", "Konversi PDF ke teks"],
        ["pandas", "\u22651.5.0", "Manipulasi data tabular"],
        ["numpy", "\u22651.23.0", "Komputasi numerik"],
        ["matplotlib", "\u22653.6.0", "Visualisasi data"],
    ],
    caption="Pustaka dan Versi yang Digunakan",
    table_num="I"
)

add_heading(doc, "B. Spesifikasi Model IndoBERT", level=2)
add_body_para(doc,
    "Model yang digunakan adalah indobenchmark/indobert-base-p1 dari Hugging Face Model Hub. Model ini memiliki "
    "arsitektur BERT-base dengan 124.441.344 parameter yang telah dilatih pada korpus bahasa Indonesia berskala besar [5]. "
    "Konfigurasi inferensi disajikan pada Tabel II."
)

add_table(doc,
    headers=["Parameter", "Nilai"],
    rows=[
        ["Model", "indobenchmark/indobert-base-p1"],
        ["Jumlah parameter", "124.441.344"],
        ["Dimensi embedding", "768"],
        ["Panjang token maksimum", "512"],
        ["Strategi pooling", "Mean Pooling"],
        ["Normalisasi", "L2 Normalization"],
        ["Ukuran batch", "8"],
        ["Perangkat komputasi", "CPU"],
    ],
    caption="Konfigurasi Model IndoBERT",
    table_num="II"
)

add_heading(doc, "C. Pipeline Implementasi", level=2)
add_body_para(doc,
    "Implementasi dilakukan secara sekuensial melalui lima notebook Jupyter: (1) 01_case_base.ipynb \u2014 konversi 35 PDF ke "
    "teks bersih; (2) 02_case_representation.ipynb \u2014 ekstraksi 18 atribut terstruktur; (3) 03_case_retrieval.ipynb \u2014 "
    "embedding IndoBERT dan pembagian 80:20; (4) 04_solution_reuse.ipynb \u2014 prediksi weighted similarity; dan "
    "(5) 05_evaluation.ipynb \u2014 perhitungan seluruh metrik evaluasi."
)

add_heading(doc, "D. Struktur Data", level=2)
add_table(doc,
    headers=["Karakteristik", "Nilai"],
    rows=[
        ["Jumlah dokumen", "35"],
        ["Rata-rata jumlah kata", "8.662 kata"],
        ["Minimum jumlah kata", "1.934 kata (case_029)"],
        ["Maksimum jumlah kata", "83.980 kata (case_010)"],
        ["Dokumen valid", "35/35 (100%)"],
        ["Rasio pembagian data", "80:20 (28 train, 7 test)"],
        ["Rentang tahun putusan", "2008\u20132025"],
    ],
    caption="Statistik Deskriptif Dataset",
    table_num="III"
)

# ==============================================================
# IV. HASIL DAN EVALUASI
# ==============================================================
add_heading(doc, "IV. Hasil dan Evaluasi")

add_heading(doc, "A. Hasil Ekstraksi Metadata", level=2)
add_body_para(doc,
    "Tingkat keberhasilan ekstraksi metadata dari 35 dokumen disajikan pada Tabel IV. Nomor perkara dan tanggal mencapai "
    "tingkat ekstraksi sempurna, sementara vonis mengalami tingkat ekstraksi paling rendah (17,1%) akibat variasi "
    "formulasi kalimat vonis dalam dokumen putusan."
)
add_table(doc,
    headers=["Atribut", "Berhasil", "Gagal", "Tingkat Keberhasilan"],
    rows=[
        ["Nomor perkara", "35", "0", "100,0%"],
        ["Tanggal", "35", "0", "100,0%"],
        ["Pasal", "34", "1", "97,1%"],
        ["Terdakwa", "30", "5", "85,7%"],
        ["Amar putusan", "32", "3", "91,4%"],
        ["Vonis", "6", "29", "17,1%"],
    ],
    caption="Tingkat Keberhasilan Ekstraksi Metadata",
    table_num="IV"
)

add_heading(doc, "B. Hasil Evaluasi Retrieval", level=2)
add_body_para(doc,
    "Evaluasi retrieval dilakukan terhadap 7 kasus testing, namun hanya 3 yang memiliki ground truth relevan di training set. "
    "Tabel V menyajikan hasil evaluasi retrieval."
)
add_table(doc,
    headers=["K", "Precision@K", "Recall@K", "F1@K", "MRR"],
    rows=[
        ["1", "0,0000", "0,0000", "0,0000", "0,0000"],
        ["3", "0,0000", "0,0000", "0,0000", "0,0000"],
        ["5", "0,0667", "0,1111", "0,0833", "0,0667"],
    ],
    caption="Metrik Evaluasi Retrieval (n=3 query dengan ground truth)",
    table_num="V"
)

add_heading(doc, "C. Hasil Evaluasi Prediksi", level=2)
add_body_para(doc,
    "Evaluasi prediksi dilakukan terhadap seluruh 7 kasus testing. Tabel VI menyajikan detail prediksi per kasus, "
    "sementara Tabel VII dan VIII merangkum metrik keseluruhan."
)
add_table(doc,
    headers=["Kasus", "Vonis Aktual", "Vonis Prediksi", "Skor Sim. Tertinggi", "Status"],
    rows=[
        ["case_014", "BEBAS", "PENJARA_LAIN", "0,7953", "Salah"],
        ["case_016", "PENJARA_LAIN", "PENJARA_LAIN", "0,8450", "Benar"],
        ["case_020", "PENJARA_LAIN", "PENJARA_LAIN", "0,8178", "Benar"],
        ["case_022", "PENJARA_LAIN", "PENJARA_LAIN", "0,8342", "Benar"],
        ["case_025", "BEBAS", "PENJARA_LAIN", "0,7620", "Salah"],
        ["case_027", "PENJARA_LAIN", "PENJARA_LAIN", "0,8505", "Benar"],
        ["case_030", "PENJARA_LAIN", "PENJARA_LAIN", "0,8600", "Benar"],
    ],
    caption="Detail Hasil Prediksi per Kasus Testing",
    table_num="VI"
)
add_table(doc,
    headers=["Metrik", "Nilai"],
    rows=[
        ["Accuracy", "0,7143 (71,43%)"],
        ["Precision (weighted)", "0,5102 (51,02%)"],
        ["Recall (weighted)", "0,7143 (71,43%)"],
        ["F1-Score (weighted)", "0,5952 (59,52%)"],
    ],
    caption="Metrik Evaluasi Prediksi Keseluruhan",
    table_num="VII"
)
add_table(doc,
    headers=["Kategori", "Precision", "Recall", "F1-Score", "Support"],
    rows=[
        ["BEBAS", "0,00", "0,00", "0,00", "2"],
        ["PENJARA_LAIN", "0,71", "1,00", "0,83", "5"],
        ["Macro average", "0,36", "0,50", "0,42", "7"],
        ["Weighted average", "0,51", "0,71", "0,60", "7"],
    ],
    caption="Classification Report per Kategori Vonis",
    table_num="VIII"
)

add_heading(doc, "D. Visualisasi Hasil", level=2)
# Insert evaluation plot if it exists
eval_plot_path = os.path.join(os.path.dirname(__file__), "data", "eval", "evaluation_plot.png")
if os.path.exists(eval_plot_path):
    doc.add_picture(eval_plot_path, width=Cm(15))
    fig_para = doc.add_paragraph(
        "Gambar 2. Visualisasi hasil evaluasi: (a) metrik retrieval per K, "
        "(b) metrik prediksi weighted similarity, (c) distribusi vonis dalam dataset"
    )
    fig_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fig_para.runs:
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.italic = True
else:
    add_body_para(doc, "[Gambar 2: Sisipkan file evaluation_plot.png dari data/eval/evaluation_plot.png]")

# ==============================================================
# V. DISKUSI
# ==============================================================
add_heading(doc, "V. Diskusi")

add_heading(doc, "A. Analisis Performa Retrieval", level=2)
add_body_para(doc,
    "Performa retrieval yang rendah (MRR = 0,0 pada K=1 dan K=3) disebabkan beberapa faktor. Pertama, dari 7 kasus testing "
    "hanya 3 yang memiliki ground truth relevan berdasarkan kecocokan pasal \u2014 pendekatan yang terlalu simplistis karena "
    "relevansi antar kasus hukum tidak hanya ditentukan oleh kesamaan pasal [10]. Kedua, pembatasan panjang teks (300 "
    "karakter per komponen) menyebabkan hilangnya informasi penting. Ketiga, ukuran training set yang kecil (28 kasus) "
    "membatasi kemungkinan menemukan kasus relevan."
)

add_heading(doc, "B. Analisis Performa Prediksi", level=2)
add_body_para(doc,
    "Akurasi 71,43% menunjukkan bahwa sistem mampu memprediksi vonis dengan tingkat kebenaran yang moderat. Namun, sistem "
    "gagal total mengidentifikasi kelas BEBAS (recall = 0,00) akibat class imbalance \u2014 hanya 2 dari 7 kasus testing "
    "berlabel BEBAS sementara 5 berlabel PENJARA_LAIN [15]. Weighted similarity voting cenderung memprediksi kategori "
    "mayoritas yang mendominasi training set."
)

add_heading(doc, "C. Analisis Kegagalan", level=2)
add_body_para(doc,
    "Dua kasus gagal diprediksi (case_014 dan case_025) keduanya merupakan kasus dengan vonis pembebasan. Case_014 "
    "(Putusan 191/PID.B/2012/PN.TRK) dibebaskan karena alasan yuridis yang tidak tercermin secara semantik dalam "
    "komponen teks embedding. Case_025 (Putusan 2418/PID.B/2005/PN.SBY) melibatkan rektor universitas yang dilepaskan "
    "dari segala tuntutan dengan skor similaritas yang lebih rendah (0,7620), mengindikasikan konteks yang cukup unik."
)

add_heading(doc, "D. Keterbatasan Tahap Revise dan Retain", level=2)
add_body_para(doc,
    "Dalam siklus CBR klasik [1], tahap Revise dan Retain merupakan komponen penting. Implementasi saat ini belum "
    "mengimplementasikan mekanisme adaptasi solusi secara otomatis (Revise) maupun penyimpanan kasus baru ke dalam "
    "case base (Retain). Penelitian selanjutnya dapat mengeksplorasi pendekatan hybrid yang mengintegrasikan umpan "
    "balik pakar hukum dalam tahap revise [3]."
)

add_heading(doc, "E. Kualitas Data dan Dampaknya", level=2)
add_body_para(doc,
    "Tingkat ekstraksi vonis yang hanya 17,1% (6 dari 35 kasus) merupakan masalah kritis yang berpotensi mengintroduksi "
    "noise signifikan dalam evaluasi. Penyebabnya adalah variasi formulasi kalimat vonis dalam putusan Mahkamah Agung "
    "yang tidak mengikuti format standar konsisten [9]."
)

add_heading(doc, "F. Rekomendasi Perbaikan", level=2)
recs = [
    "Perbaikan preprocessing: Memperkaya pola regex dengan variasi formulasi seperti 'dipidana', 'dijatuhi pidana', dan 'dihukum penjara'.",
    "Augmentasi dataset: Meningkatkan jumlah dokumen menjadi minimal 100 untuk distribusi kelas yang lebih seimbang.",
    "Fine-tuning model: Melakukan fine-tuning IndoBERT pada korpus dokumen hukum Indonesia [14].",
    "Hybrid retrieval: Mengkombinasikan IndoBERT embedding dengan BM25 untuk meningkatkan akurasi retrieval [13].",
    "Constraint-based retrieval: Menambahkan hard constraint berupa kecocokan pasal sebagai filter awal.",
    "Threshold confidence: Menerapkan ambang batas kepercayaan untuk menandai prediksi yang tidak yakin.",
]
for i, rec in enumerate(recs, 1):
    para = doc.add_paragraph()
    para_fmt(para)
    run = para.add_run(f"{i}. {rec}")
    set_font(run)

# ==============================================================
# VI. KESIMPULAN
# ==============================================================
add_heading(doc, "VI. Kesimpulan")

add_body_para(doc,
    "Penelitian ini telah berhasil mengimplementasikan sistem Case-Based Reasoning (CBR) berbasis IndoBERT untuk analisis "
    "putusan pidana pemalsuan dari Mahkamah Agung Republik Indonesia. Sistem dikembangkan melalui lima tahap: pembangunan "
    "case base dari 35 dokumen PDF, representasi kasus terstruktur melalui ekstraksi regex, case retrieval menggunakan "
    "embedding IndoBERT 768 dimensi dengan cosine similarity, prediksi vonis melalui weighted similarity voting, dan "
    "evaluasi komprehensif."
)
add_body_para(doc,
    "Hasil evaluasi menunjukkan akurasi prediksi 71,43% dengan F1-Score 0,5952 (weighted). Sistem mampu memprediksi vonis "
    "penjara dengan recall sempurna (1,00) namun gagal mengidentifikasi kasus pembebasan (recall = 0,00 untuk BEBAS). "
    "Performa retrieval masih rendah (MRR = 0,0667 pada K=5) dipengaruhi keterbatasan ground truth dan ukuran dataset."
)
add_body_para(doc,
    "Keterbatasan utama meliputi: (1) ukuran dataset kecil (35 dokumen), (2) rendahnya tingkat ekstraksi vonis (17,1%), "
    "(3) ketidakseimbangan distribusi kelas, dan (4) belum sepenuhnya mengimplementasikan tahap Revise dan Retain. "
    "Penelitian selanjutnya disarankan memperbesar dataset, melakukan fine-tuning IndoBERT pada domain hukum, "
    "mengimplementasikan hybrid retrieval (BM25 + embedding), serta mengeksplorasi adaptasi solusi otomatis."
)

# ==============================================================
# DAFTAR PUSTAKA
# ==============================================================
add_heading(doc, "Daftar Pustaka")

references = [
    "[1] A. Aamodt and E. Plaza, \"Case-Based Reasoning: Foundational Issues, Methodological Variations, and System Approaches,\" AI Communications, vol. 7, no. 1, pp. 39\u201359, 1994.",
    "[2] J. L. Kolodner, Case-Based Reasoning. San Mateo, CA: Morgan Kaufmann, 1993.",
    "[3] R. L. de M\u00e1ntaras, D. McSherry, D. Bridge, D. Leake, B. Smyth, S. Craw, B. Faltings, M. L. Maher, M. T. Cox, K. Forbus, M. Keane, A. Aamodt, and I. Watson, \"Retrieval, Reuse, Revision, and Retention in Case-Based Reasoning,\" The Knowledge Engineering Review, vol. 20, no. 3, pp. 215\u2013240, 2005.",
    "[4] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, \"BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding,\" in Proc. of NAACL-HLT, Minneapolis, MN, 2019, pp. 4171\u20134186.",
    "[5] B. Wilie, K. Vincentio, G. I. Winata, S. Cahyawijaya, X. Li, Z. Y. Lim, S. Soleman, R. Mahendra, P. Fung, S. Bahar, and A. Purwarianti, \"IndoNLU: Benchmark and Resources for Evaluating Indonesian Natural Language Understanding,\" in Proc. of AACL-IJCNLP, Suzhou, China, 2020, pp. 843\u2013857.",
    "[6] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,\" in Proc. of EMNLP-IJCNLP, Hong Kong, China, 2019, pp. 3982\u20133992.",
    "[7] F. Koto, A. Rahimi, J. H. Lau, and T. Baldwin, \"IndoLEM and IndoBERT: A Benchmark Dataset and Pre-trained Language Model for Indonesian NLP,\" in Proc. of COLING, Barcelona, Spain, 2020, pp. 757\u2013770.",
    "[8] J. Collenette, K. Atkinson, and T. Bench-Capon, \"Explainable AI Tools for Legal Reasoning about Cases: A Study on the European Court of Human Rights,\" Artificial Intelligence, vol. 317, pp. 103861, 2023.",
    "[9] J. Cui, X. Shen, and S. Wen, \"A Survey on Legal Judgment Prediction: Datasets, Metrics, Models and Challenges,\" IEEE Access, vol. 11, pp. 102050\u2013102071, 2023.",
    "[10] F. Ariai, J. Mackenzie, and G. Demartini, \"Natural Language Processing for the Legal Domain: A Survey of Tasks, Datasets, Models, and Challenges,\" ACM Computing Surveys, vol. 58, no. 6, pp. 1\u201337, 2025.",
    "[11] A. P. Kemala and H. A. Shiddiqi, \"Analysis of Indonesian Language Dataset for Tax Court Cases: Multiclass Classification of Court Verdicts,\" Jurnal Riset Informatika, vol. 5, no. 3, pp. 419\u2013424, 2023.",
    "[12] E. Yulianti, N. Bhary, J. Abdurrohman, F. W. Dwitilas, E. Q. Nuranti, and H. S. Husin, \"Named Entity Recognition on Indonesian Legal Documents: A Dataset and Study Using Transformer-Based Models,\" International Journal of Electrical and Computer Engineering (IJECE), vol. 14, no. 5, pp. 5489\u20135501, 2024.",
    "[13] C. D. Manning, P. Raghavan, and H. Sch\u00fctze, Introduction to Information Retrieval. Cambridge: Cambridge University Press, 2008.",
    "[14] M. A. Ibrahim, A. T. Handoyo, and M. S. Anggreainy, \"Hybrid Deep Learning for Legal Text Analysis: Predicting Punishment Durations in Indonesian Court Rulings,\" arXiv preprint, arXiv:2410.20104, 2024.",
    "[15] H. He and E. A. Garcia, \"Learning from Imbalanced Data,\" IEEE Transactions on Knowledge and Data Engineering, vol. 21, no. 9, pp. 1263\u20131284, 2009.",
]

for ref in references:
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.first_line_indent = Pt(-18)
    para.paragraph_format.left_indent = Pt(18)
    run = para.add_run(ref)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)

# ==============================================================
# SAVE
# ==============================================================
output_path = os.path.join(os.path.dirname(__file__), "artikel_cbr.docx")
doc.save(output_path)
print(f"[OK] File berhasil disimpan: {output_path}")
